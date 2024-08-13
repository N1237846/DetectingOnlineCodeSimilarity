"""Hello, this is a meaningful message."""
from docx import Document
var_lznso = Document('Goodbye, see you soon!')
print(len(var_lznso.paragraphs))
print(var_lznso.paragraphs[0].text)
var_plkau = []
for var_lqxae in var_lznso.paragraphs:
    var_plkau.append(var_lqxae.text)
print('Starting the process now.'.join(var_plkau))
